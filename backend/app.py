"""
School Management System - Backend API
Flask API for report card generation using Word templates
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import json
import requests
from datetime import datetime
import logging
from docx import Document
from docx.shared import Inches
import io
import base64

# Firebase Admin SDK
try:
    from firebase_admin import credentials, initialize_app, firestore
    FIREBASE_AVAILABLE = True
except ImportError:
    FIREBASE_AVAILABLE = False
    print("⚠️ Firebase Admin SDK not available. Install with: pip install firebase-admin")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)  # Enable CORS for frontend integration

# Configuration
FIREBASE_PROJECT_ID = os.getenv('FIREBASE_PROJECT_ID', 'ass-sms')
FIREBASE_API_KEY = os.getenv('FIREBASE_API_KEY', 'AIzaSyC13_vM3RlSSxCDXTiKafekCBPpejtSGWc')

# Initialize Firebase Admin SDK
def initialize_firebase():
    """Initialize Firebase Admin SDK with service account"""
    if not FIREBASE_AVAILABLE:
        logger.warning("Firebase Admin SDK not available")
        return None
    
    try:
        # Get credentials from environment variable
        creds_json = os.getenv('FIREBASE_CREDENTIALS_JSON')
        if creds_json:
            creds_dict = json.loads(creds_json)
            cred = credentials.Certificate(creds_dict)
            initialize_app(cred)
            logger.info("✅ Firebase Admin SDK initialized successfully")
            return firestore.client()
        else:
            logger.warning("❌ FIREBASE_CREDENTIALS_JSON not found in environment variables")
            return None
    except Exception as e:
        logger.error(f"❌ Firebase initialization failed: {e}")
        return None

# Initialize Firebase
db = initialize_firebase()

class ReportCardGenerator:
    """Handles report card generation using Word templates"""
    
    def __init__(self):
        self.template_path = os.path.join(os.path.dirname(__file__), 'report_template.docx')
        
    def populate_word_template(self, student_data, grades_data, term, class_name):
        """Populate Word template with student data"""
        try:
            # Load the Word template
            doc = Document(self.template_path)
            
            # Replace placeholders in the document
            self._replace_placeholders(doc, student_data, grades_data, term, class_name)
            
            # Save to memory
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output
            
        except Exception as e:
            logger.error(f"Word template population failed: {str(e)}")
            raise e
    
    def _replace_placeholders(self, doc, student_data, grades_data, term, class_name):
        """Replace placeholders in Word document"""
        # School information
        replacements = {
            '{SchoolName}': 'ADILANG SECONDARY SCHOOL',
            '{SchoolAddress}': 'P.O.BOX 13, PADER-AGAGO DISTRICT',
            '{SchoolPhone}': 'TEL: 0773221580/0782634466/0782446279/0770685882',
            '{ReportTitle}': f'END OF TERM {term.upper()} ASSESSMENT REPORT CARD 2025',
            '{Term}': term,
            '{Year}': '2025',
            
            # Student information
            '{StudentName}': student_data.get('personalInfo', {}).get('fullName', ''),
            '{StudentClass}': class_name,
            '{StudentStream}': student_data.get('class', ''),
            '{StudentLIN}': student_data.get('admissionNumber', ''),
            '{PaymentCode}': student_data.get('paymentCode', ''),
            '{StudentSex}': student_data.get('sex', ''),
        }
        
        # Calculate academic data
        total_marks = sum(grade.get('finalGrade', {}).get('percentage', 0) for grade in grades_data)
        average_score = total_marks / len(grades_data) if grades_data else 0
        
        replacements.update({
            '{TotalMarks}': str(total_marks),
            '{AverageScore}': f"{average_score:.1f}",
            '{OverallPerformance}': self._get_performance_grade(average_score)
        })
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    def _get_performance_grade(self, average_score):
        """Convert average score to letter grade"""
        if average_score >= 80:
            return 'A'
        elif average_score >= 70:
            return 'B'
        elif average_score >= 60:
            return 'C'
        elif average_score >= 50:
            return 'D'
        else:
            return 'F'

# Initialize report card generator
report_generator = ReportCardGenerator()

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'School Management System API'
    })

@app.route('/api/generate-word-report', methods=['POST'])
def generate_word_report():
    """Generate Word document report card"""
    try:
        data = request.get_json()
        
        student_id = data['studentId']
        term = data['term']
        class_name = data['className']
        
        # Get data
        student_data = get_student_data(student_id)
        grades_data = get_grades_data(student_id, term)
        
        if not student_data:
            return jsonify({
                'success': False,
                'message': 'Student not found'
            }), 404
        
        # Generate Word document
        word_doc = report_generator.populate_word_template(
            student_data, grades_data, term, class_name
        )
        
        # Return as base64 encoded string
        word_content = word_doc.getvalue()
        word_base64 = base64.b64encode(word_content).decode('utf-8')
        
        return jsonify({
            'success': True,
            'message': 'Word document generated successfully',
            'data': {
                'content': word_base64,
                'filename': f'Report_Card_{student_data.get("personalInfo", {}).get("fullName", "Unknown")}_{term}.docx'
            }
        })
        
    except Exception as e:
        logger.error(f"Word report generation failed: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'Internal server error',
            'error': str(e)
        }), 500

@app.route('/api/batch-generate', methods=['POST'])
def batch_generate():
    """Generate report cards for multiple students"""
    try:
        data = request.get_json()
        
        students = data.get('students', [])
        term = data.get('term')
        class_name = data.get('class')
        
        if not students or not term or not class_name:
            return jsonify({
                'success': False,
                'message': 'Missing required fields: students, term, class'
            }), 400
        
        results = []
        for student_id in students:
            try:
                student_data = get_student_data(student_id)
                if student_data:
                    grades_data = get_grades_data(student_id, term)
                    word_doc = report_generator.populate_word_template(
                        student_data, grades_data, term, class_name
                    )
                    
                    # Save document (you can implement file saving logic here)
                    results.append({
                        'studentId': student_id,
                        'success': True,
                        'message': 'Report card generated successfully'
                    })
                else:
                    results.append({
                        'studentId': student_id,
                        'success': False,
                        'message': 'Student not found'
                    })
            except Exception as e:
                results.append({
                    'studentId': student_id,
                    'success': False,
                    'message': str(e)
                })
        
        return jsonify({
            'success': True,
            'message': f'Batch processing completed for {len(students)} students',
            'results': results
        })
        
    except Exception as e:
        logger.error(f"Batch generation failed: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'Internal server error',
            'error': str(e)
        }), 500

def get_student_data(student_id):
    """Get student data from Firestore (implement based on your Firebase setup)"""
    # This is a placeholder - implement based on your Firebase configuration
    # You can use the firebase-admin SDK or REST API
    try:
        # Example implementation using REST API
        url = f"https://firestore.googleapis.com/v1/projects/{FIREBASE_PROJECT_ID}/databases/(default)/documents/students/{student_id}"
        response = requests.get(url, params={'key': FIREBASE_API_KEY})
        
        if response.status_code == 200:
            data = response.json()
            return data.get('fields', {})
        else:
            logger.error(f"Failed to fetch student data: {response.status_code}")
            return None
            
    except Exception as e:
        logger.error(f"Error fetching student data: {str(e)}")
        return None

def get_grades_data(student_id, term):
    """Get grades data from Firestore"""
    try:
        # Example implementation using REST API
        url = f"https://firestore.googleapis.com/v1/projects/{FIREBASE_PROJECT_ID}/databases/(default)/documents:runQuery"
        
        query = {
            "structuredQuery": {
                "from": [{"collectionId": "grades"}],
                "where": {
                    "compositeFilter": {
                        "op": "AND",
                        "filters": [
                            {
                                "fieldFilter": {
                                    "field": {"fieldPath": "studentId"},
                                    "op": "EQUAL",
                                    "value": {"stringValue": student_id}
                                }
                            },
                            {
                                "fieldFilter": {
                                    "field": {"fieldPath": "term"},
                                    "op": "EQUAL",
                                    "value": {"stringValue": term}
                                }
                            }
                        ]
                    }
                }
            }
        }
        
        response = requests.post(url, json=query, params={'key': FIREBASE_API_KEY})
        
        if response.status_code == 200:
            data = response.json()
            return [doc.get('document', {}).get('fields', {}) for doc in data.get('documents', [])]
        else:
            logger.error(f"Failed to fetch grades data: {response.status_code}")
            return []
            
    except Exception as e:
        logger.error(f"Error fetching grades data: {str(e)}")
        return []

@app.route('/api/test-firebase', methods=['GET'])
def test_firebase():
    """Test Firebase connection"""
    if not db:
        return jsonify({
            'success': False,
            'message': 'Firebase not initialized. Check FIREBASE_CREDENTIALS_JSON environment variable.'
        }), 500
    
    try:
        # Test Firestore connection
        doc_ref = db.collection('test').document('connection')
        doc_ref.set({
            'message': 'Firebase connection test',
            'timestamp': datetime.now().isoformat(),
            'project_id': FIREBASE_PROJECT_ID
        })
        
        # Read it back
        doc = doc_ref.get()
        if doc.exists:
            return jsonify({
                'success': True,
                'message': 'Firebase connection successful',
                'data': doc.to_dict(),
                'project_id': FIREBASE_PROJECT_ID
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Firebase connection failed - document not found'
            })
    except Exception as e:
        logger.error(f"Firebase test error: {e}")
        return jsonify({
            'success': False,
            'message': f'Firebase error: {str(e)}'
        }), 500

@app.route('/api/test', methods=['GET'])
def test_endpoint():
    """Test endpoint for development"""
    return jsonify({
        'message': 'Backend API is working!',
        'timestamp': datetime.now().isoformat(),
        'template_path': report_generator.template_path,
        'template_exists': os.path.exists(report_generator.template_path),
        'firebase_available': FIREBASE_AVAILABLE,
        'firebase_initialized': db is not None,
        'project_id': FIREBASE_PROJECT_ID
    })

if __name__ == '__main__':
    # Get port from environment variable (for Render) or default to 5000
    port = int(os.getenv('PORT', 5000))
    debug = os.getenv('FLASK_DEBUG', 'False').lower() == 'true'
    
    # Production server for Render
    app.run(debug=debug, host='0.0.0.0', port=port)
